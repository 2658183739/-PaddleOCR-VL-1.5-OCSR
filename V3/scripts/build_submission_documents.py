#!/usr/bin/env python3
"""Build the final Chinese DOCX reports for the OCSR V3 submission."""

from __future__ import annotations

import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


V3_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = V3_ROOT.parents[1]
FINAL_DIR = WORKSPACE / "决赛"
ASSET_DIR = V3_ROOT / "presentation" / "assets"
BUILD_DIR = ASSET_DIR / "report_generated"
FINAL_RESULTS = V3_ROOT / "evidence" / "FINAL_RESULTS.json"

NAVY = "17212D"
INK = "182230"
TEAL = "087F72"
BLUE = "2458A6"
RED = "B42318"
AMBER = "B76A00"
MUTED = "687386"
LINE = "D9DEE7"
SOFT = "F3F6F8"
WHITE = "FFFFFF"

FONT_PATH = Path("C:/Windows/Fonts/msyh.ttc")
FONT_BOLD_PATH = Path("C:/Windows/Fonts/msyhbd.ttc")
def image_font(size, bold=False):
    target = FONT_BOLD_PATH if bold and FONT_BOLD_PATH.exists() else FONT_PATH
    if target.exists():
        return ImageFont.truetype(str(target), size=size)
    return ImageFont.load_default()


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def set_cell_text(cell, text, *, color=INK, bold=False, size=8.5, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].height = Cm(0.78)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, color=WHITE, bold=True, size=8.2)
        if widths:
            cell.width = Cm(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, size=font_size)
            if widths:
                cells[index].width = Cm(widths[index])
            if row_index % 2:
                set_cell_shading(cells[index], SOFT)
            set_cell_border(cells[index], bottom={"val": "single", "sz": "4", "color": LINE})
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_paragraph(doc, text, *, bold_lead=None, color=INK, size=10, space_after=7, keep=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True, color=TEAL)
        run = paragraph.add_run(text[len(bold_lead):])
        set_run_font(run, size=size, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color)
    return paragraph


def add_callout(doc, title, text, *, color=TEAL):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(0.22)
    table.columns[1].width = Cm(16.6)
    set_cell_shading(table.cell(0, 0), color)
    set_cell_shading(table.cell(0, 1), "F4F7F9")
    cell = table.cell(0, 1)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=9.2, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    run = p2.add_run(text)
    set_run_font(run, size=8.7, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_heading(doc, text, level=1, *, new_page=False):
    if new_page:
        doc.add_page_break()
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 4)
    return paragraph


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=8, color=MUTED)


def add_picture(doc, path: Path, caption: str, width=6.55):
    if not path.exists():
        add_callout(doc, "缺失图像", f"构建时未找到 {path.name}，正文事实不受影响。", color=RED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def configure_document(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.05)
    section.bottom_margin = Cm(1.75)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for level, size, color in [(1, 18, NAVY), (2, 13, TEAL), (3, 10.5, BLUE)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(short_title + "  |  PaddleOCR-VL OCSR V3")
    set_run_font(run, size=7.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("V3 final  ·  2026-07-20  ·  第 ")
    set_run_font(run, size=7.5, color=MUTED)
    add_field(p, "PAGE")
    run = p.add_run(" 页")
    set_run_font(run, size=7.5, color=MUTED)


def add_cover(doc: Document, title: str, subtitle: str, report_type: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17)
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    cell.height = Cm(3.1)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("PADDLEOCR-VL · OCSR · V3 FINAL")
    set_run_font(run, size=10, bold=True, color="8DD8CD")
    p2 = cell.add_paragraph()
    run = p2.add_run(report_type)
    set_run_font(run, size=9, color=WHITE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    set_run_font(run, size=28, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run(subtitle)
    set_run_font(run, size=13, color=TEAL)

    add_table(
        doc,
        ["训练 strict", "开发面板", "Locked wild", "Scaffold novel", "最终解码"],
        [["22,762", "753 + 754", "301 / 62 papers", "134", "Beam4 / Return4"]],
        widths=[3.25, 3.25, 3.65, 3.25, 3.55],
        font_size=9.3,
    )
    add_callout(
        doc,
        "一句话结论",
        "V3 用可追溯的数据角色、按 molecule/paper group 防泄漏、受控 LoRA continuation 与一次性 locked test，证明模型已具备 OCSR 能力，同时诚实暴露论文真实域仍是主要瓶颈。",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(62)
    run = p.add_run("版本：V3 final\n日期：2026-07-20\n公开地址：GitHub + Hugging Face（revision 固定）")
    set_run_font(run, size=9, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document, sections):
    add_heading(doc, "阅读导航", level=1)
    add_paragraph(doc, "本页先给出材料结构。Word 打开文档后可右键目录并选择“更新域”，生成带页码的自动目录。", color=MUTED, size=9)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    table_rows = [(str(i + 1).zfill(2), title, purpose) for i, (title, purpose) in enumerate(sections)]
    add_table(doc, ["章节", "内容", "回答的问题"], table_rows, widths=[1.5, 5.4, 10.4], font_size=8.4)
    doc.add_page_break()


def save_figure(image, name):
    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    path = BUILD_DIR / name
    image.save(path, format="PNG", optimize=True)
    return path


def build_figures():
    mixture = [
        ("USPTO", 5043, TEAL), ("UOB", 4869, BLUE), ("real-world", 4329, AMBER),
        ("MolGrapher synthetic", 4000, "6F7785"), ("30K clean", 1501, "5B8C5A"),
        ("30K abbreviated", 1507, "9B5DA5"), ("30K large", 1513, RED),
    ]
    image = Image.new("RGB", (2200, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "V3 final control：七部分训练混合", font=image_font(54, True), fill="#17212D")
    draw.text((80, 123), "训练记录数包含 repeat/cap 后的采样权重，不等于独立图片数", font=image_font(28), fill="#687386")
    max_count = 5600
    bar_left, bar_width = 620, 1230
    for index, (label, count, color) in enumerate(mixture):
        y = 205 + index * 98
        draw.text((80, y + 14), label, font=image_font(31, True), fill="#344055")
        draw.rounded_rectangle((bar_left, y, bar_left + bar_width, y + 58), radius=8, fill="#EEF1F4")
        length = int(bar_width * count / max_count)
        draw.rounded_rectangle((bar_left, y, bar_left + length, y + 58), radius=8, fill="#" + color)
        draw.text((bar_left + bar_width + 35, y + 10), f"{count:,}  ·  {count/22762:.1%}", font=image_font(28, True), fill="#182230")
    mixture_path = save_figure(image, "training-mixture.png")

    panels = [("core dev", 753, BLUE), ("region dev", 754, TEAL), ("locked wild", 301, RED), ("scaffold novel", 134, AMBER), ("symbolic", 460, "6F7785")]
    image = Image.new("RGB", (2200, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "评测数据按角色分层，而不是混成一个总分", font=image_font(54, True), fill="#17212D")
    draw.text((80, 123), "Development 负责选择；locked 只负责一次性外推报告；symbolic 独立计分", font=image_font(28), fill="#687386")
    baseline_y, chart_top = 730, 240
    for grid in (0, 250, 500, 750):
        y = baseline_y - int((baseline_y - chart_top) * grid / 800)
        draw.line((110, y, 2100, y), fill="#E5E8ED", width=2)
        draw.text((35, y - 18), str(grid), font=image_font(24), fill="#687386")
    for index, (label, count, color) in enumerate(panels):
        x0 = 210 + index * 390
        height = int((baseline_y - chart_top) * count / 800)
        draw.rounded_rectangle((x0, baseline_y - height, x0 + 220, baseline_y), radius=10, fill="#" + color)
        draw.text((x0 + 110, baseline_y - height - 55), str(count), anchor="mm", font=image_font(34, True), fill="#182230")
        draw.text((x0 + 110, baseline_y + 55), label, anchor="mm", font=image_font(27, True), fill="#344055")
    eval_path = save_figure(image, "evaluation-roles.png")

    labels = ["V2-1 core\nbaseline", "V2-1 core\nstable", "V2-1 region\nbaseline", "V2-1 region\nstable", "V3 ckpt-1400\ngreedy", "V3 ckpt-1400\nbeam4", "locked wild", "scaffold novel"]
    values = [37.03, 45.89, 38.44, 43.77, 35.97, 42.07, 22.92, 13.43]
    colors = ["#6F7785", "#087F72", "#6F7785", "#087F72", "#2458A6", "#087F72", "#B42318", "#B76A00"]
    image = Image.new("RGB", (2400, 980), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "同口径内部比较；跨角色结果不得纵向混算", font=image_font(54, True), fill="#17212D")
    draw.text((80, 123), "Canonical exact (%)", font=image_font(28), fill="#687386")
    baseline_y, chart_top = 750, 220
    for grid in (0, 10, 20, 30, 40, 50):
        y = baseline_y - int((baseline_y - chart_top) * grid / 52)
        draw.line((100, y, 2320, y), fill="#E5E8ED", width=2)
        draw.text((38, y - 15), str(grid), font=image_font(22), fill="#687386")
    for index, (label, value, color) in enumerate(zip(labels, values, colors)):
        x0 = 150 + index * 277
        height = int((baseline_y - chart_top) * value / 52)
        draw.rounded_rectangle((x0, baseline_y - height, x0 + 170, baseline_y), radius=8, fill=color)
        draw.text((x0 + 85, baseline_y - height - 42), f"{value:.2f}", anchor="mm", font=image_font(27, True), fill="#182230")
        lines = label.split("\n")
        draw.text((x0 + 85, baseline_y + 47), lines[0], anchor="mm", font=image_font(22, True), fill="#344055")
        if len(lines) > 1:
            draw.text((x0 + 85, baseline_y + 80), lines[1], anchor="mm", font=image_font(21), fill="#687386")
    results_path = save_figure(image, "results-boundaries.png")

    steps = ["输入契约", "视觉质检", "候选生成", "RDKit 校验", "一致性排序", "证据汇总"]
    image = Image.new("RGB", (2400, 650), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), "MolTrace OCSR Agent：可审计的六步决策链", font=image_font(54, True), fill="#17212D")
    for index, label in enumerate(steps):
        x0 = 70 + index * 390
        draw.rounded_rectangle((x0, 230, x0 + 320, 475), radius=14, fill="#F3F6F8", outline="#D9DEE7", width=3)
        draw.text((x0 + 30, 263), f"{index+1:02d}", font=image_font(27, True), fill="#087F72")
        draw.text((x0 + 30, 345), label, font=image_font(36, True), fill="#182230")
        if index < len(steps) - 1:
            draw.line((x0 + 324, 352, x0 + 375, 352), fill="#087F72", width=6)
            draw.polygon([(x0 + 375, 352), (x0 + 353, 338), (x0 + 353, 366)], fill="#087F72")
    agent_path = save_figure(image, "agent-flow.png")
    return {"mixture": mixture_path, "eval": eval_path, "results": results_path, "agent": agent_path}


def build_main_report(figures):
    doc = Document()
    configure_document(doc, "最终版数据构建报告")
    add_cover(doc, "PaddleOCR-VL OCSR V3\n最终版数据构建报告", "从训练数据、评测集到训练、后训练与 Agent 交付的完整证据链", "FINAL DATA CONSTRUCTION & TRAINING REPORT")
    sections = [
        ("项目定义与结论", "任务究竟是什么，哪些结果可信"),
        ("V2-1 基础与 V3 增量", "为什么从已有 OCSR 能力继续训练"),
        ("数据治理与标签契约", "什么样的数据才允许进入主任务"),
        ("七部分训练数据构建", "各数据集比例如何形成、为何不平均"),
        ("评测集构建与防泄漏", "development、locked、symbolic 如何分工"),
        ("训练与消融设计", "如何避免把偶然性写成最优"),
        ("后训练与解码选择", "为什么采用 beam，拒绝 hard replay/rerank"),
        ("最终结果与误差边界", "开发结果和域外测试如何解释"),
        ("MolTrace Agent 前后端", "产品如何把模型决策变得可用、可审计"),
        ("许可、复现与交付", "公开什么、不公开什么、如何验收"),
        ("限制与后续计划", "院士最可能追问的问题"),
    ]
    add_toc(doc, sections)

    add_heading(doc, "1. 项目定义与结论", level=1)
    add_paragraph(doc, "本项目把 PaddleOCR-VL-1.5 适配为光学化学结构识别（OCSR）模型：输入一张分子结构图，输出一行可由 RDKit 解析并 canonicalize 的单分子 SMILES。主任务不接受反应式、盐/溶剂多片段、R-group symbolic 或说明性文字；这些边界样本单独进入 symbolic 轨道，防止指标被不同任务定义污染。")
    add_callout(doc, "最重要的事实", "训练 strict control 为 22,762 条；可调 development 为 753 + 754 条；最终 wild test 为 301 个唯一分子、来自 62 篇完整留出论文。V3 locked wild exact 为 22.92%、valid 为 84.72%，说明模型已会生成大量化学上有效的字符串，但真实论文图的结构完全恢复仍是核心难点。")
    add_table(doc, ["证据层", "回答的问题", "当前结论", "不可越过的边界"], [
        ["V2-1 历史面板", "模型是否学会 OCSR 格式", "core 37.03% → 45.89%；region 38.44% → 43.77%", "历史面板反复调参，不是 locked test"],
        ["V3 development", "checkpoint 与 decoder 怎么选", "ckpt-1400 + beam4/return4，macro exact 42.07%", "只回答当前 development 分布"],
        ["V3 locked wild", "论文外推如何", "exact 22.92%，valid 84.72%", "冻结后只运行一次，不回流"],
        ["Scaffold novel", "未见骨架如何", "exact 13.43%，valid 75.37%", "是 wild 子集，不是独立调参集"],
    ], widths=[3, 4.2, 5.3, 4.7])
    add_picture(doc, figures["results"], "图 1  结果只能在同一数据角色和口径内部比较")

    add_heading(doc, "2. V2-1 基础与 V3 增量", level=1, new_page=True)
    add_paragraph(doc, "V2-1 的关键贡献不是最终域外分数，而是完成了任务适配。原始 PaddleOCR-VL-1.5 在固定预算 warm-start 对照的两个 development 面板上 exact 均为 0；V2-1 continuation 的宏平均 exact 为 34.17%。另一组历史诊断中，原始模型的 canonical exact、valid、token micro-F1 和 mean Tanimoto 分别为 0.00%、30.78%、6.59% 和 0.0027，而 V2-1 LoRA export 分别为 33.77%、75.84%、70.18% 和 0.6849。")
    add_paragraph(doc, "V3 因此不从基础权重重新学习输出格式，而是从 V2-1 export 做低学习率 LoRA continuation。有限 H800 预算被用于数据因素、seed、checkpoint、hard replay 和 decoder 的受控对照。这个决策符合迁移学习常识：当目标格式能力已经建立，重复 warm-start 会浪费预算并扩大不确定性。")
    add_table(doc, ["版本/阶段", "数据角色", "核心努力", "结论"], [
        ["原始 1.5", "固定预算对照", "未做 OCSR task adaptation", "exact 0；不能直接承担 canonical 输出"],
        ["V2-1 LoRA", "历史 development", "完成任务格式适配、历史后处理和真实场景诊断", "证明继续训练路线有合理起点"],
        ["V3 probe", "legacy development", "2×2 数据因素 × 2 seed + warm-start", "control 均值最高，但不宣称普适最优"],
        ["V3 final", "development + locked", "1400-step continuation、decoder 冻结、一次性终测", "得到可审计最终证据"],
    ], widths=[3.2, 3.4, 6.1, 4.5])
    add_callout(doc, "为什么报告必须包含 V2-1", "如果只展示 V3，会看不出任务能力从哪里来；如果把 V2-1 历史面板当最终测试，又会夸大泛化。正确写法是把 V2-1 作为 task adaptation 与方法演进证据，把 V3 locked wild 作为域外结果。", color=BLUE)

    add_heading(doc, "3. 数据治理与标签契约", level=1, new_page=True)
    add_paragraph(doc, "所有正式主任务记录必须满足同一契约：一张可打开图片、一个明确目标、一个单分子 canonical SMILES。构建顺序为路径与图像检查、标签清洗、RDKit 解析、canonicalization 幂等检查、多片段与 dummy atom 拒绝、跨 split 去重、来源与难度字段记录。任务定义先冻结，训练策略不能反过来改变标签含义。")
    add_table(doc, ["门", "自动规则", "拒绝/分流对象", "原因"], [
        ["图像门", "文件存在、可打开、尺寸非零", "坏图、路径漂移、空白文件", "避免训练和评测运行时失败"],
        ["化学门", "RDKit 可解析、canonical 幂等、无 dummy atom", "非法 SMILES、占位原子", "保证 exact 的目标空间一致"],
        ["单分子门", "不含点号分隔的多个片段", "盐、溶剂、混合物", "主任务定义为单分子"],
        ["任务门", "仅 canonical SMILES", "反应式、R-group、chemfig、symbolic", "独立轨道，不能混入主分母"],
        ["泄漏门", "molecule/structure_id/paper_group 反向检查", "与 development/locked 重叠", "阻止同结构或同论文穿越 split"],
        ["许可门", "source、上游 ID、license/source URL 可追踪", "来源不明且需公开分发的数据", "代码开源不等于数据可再分发"],
    ], widths=[2.3, 5.5, 4.7, 4.7])
    add_paragraph(doc, "V2-1 clean weighted 输入为 23,047 条。V3 删除 273 条多片段、symbolic 或不收敛标签，再删除 12 条与新 held-out 分子重叠的记录，最终 strict control 为 22,762 条。repeat/cap 只改变训练抽样权重，不增加独立样本数；报告所有比例时必须明确分母是训练记录还是唯一图/唯一分子。")

    add_heading(doc, "4. 七部分训练数据构建", level=1, new_page=True)
    add_picture(doc, figures["mixture"], "图 2  七部分训练混合及科学配比")
    add_table(doc, ["部分", "记录数", "占比", "角色与清洗口径"], [
        ["USPTO", "5,043", "22.16%", "公开专利风格锚点；canonical 化、坏图/多片段剔除"],
        ["UOB", "4,869", "21.39%", "公开 OCSR benchmark 风格；标签、路径和重复检查"],
        ["real-world", "4,329", "19.02%", "拍照/扫描/页面/手写场景；只接受已有可信 SMILES"],
        ["MolGrapher synthetic", "4,000", "17.57%", "复杂结构与视觉扰动训练；不计作真实评测"],
        ["USPTO-30K clean", "1,501", "6.59%", "干净专利图；cap 约 1,500，避免压过目标域"],
        ["USPTO-30K abbreviated", "1,507", "6.62%", "缩写长尾；独立 cap"],
        ["USPTO-30K large", "1,513", "6.65%", "大图/长分子长尾；独立 cap"],
        ["合计", "22,762", "100.00%", "Final A control；记录数含训练采样权重"],
    ], widths=[4, 2.3, 2.2, 8.7])
    add_heading(doc, "4.1 为什么不平均配比", level=2)
    add_paragraph(doc, "USPTO 与 UOB 合计约 43.5%，作为 printed 能力锚点；real-world 约 19.0%，让稀缺拍照、扫描和页面退化在 batch 中持续出现；MolGrapher synthetic 约 17.6%，覆盖复杂拓扑和受控视觉扰动；三个 USPTO-30K 长尾子集各限制约 1,500 条，防止数量较大的干净专利图主导梯度。")
    add_paragraph(doc, "这个比例来自标签可信度、V2-1 错误分层、训练预算和 2×2 探索共同约束，不是凭直觉平均，也不是充分搜索后的全局最优。两个 seed 只能降低偶然性，不能支持“统计显著最优”或“适用于所有部署域”的结论。下一轮 confirmatory 应至少使用四个 seed，并用分块随机或平衡 Latin-square 控制运行顺序。")
    add_heading(doc, "4.2 五个受控混合", level=2)
    add_table(doc, ["混合", "总记录", "Strict wild", "离线退化", "实验问题"], [
        ["A control", "22,762", "0", "0", "稳定基础混合"],
        ["D wild-only", "23,562", "800", "0", "真实论文图是否有主效应"],
        ["E aug-only", "23,689", "0", "927", "离线退化是否有主效应"],
        ["B wild + aug", "24,489", "800", "927", "二者是否存在交互"],
        ["C real-heavy", "25,416", "800", "1,854", "增强剂量响应"],
    ], widths=[3.3, 2.5, 2.8, 2.8, 5.8])

    add_heading(doc, "5. 评测集构建与防泄漏", level=1, new_page=True)
    add_picture(doc, figures["eval"], "图 3  评测集按角色拆分：选模、锁定终测和 symbolic 不混算")
    add_table(doc, ["面板", "N", "独立单位", "角色", "使用规则"], [
        ["dev_legacy_core_strict", "753", "molecule / structure_id", "Core development", "可选 checkpoint、数据因素、decoder"],
        ["dev_legacy_region_strict", "754", "molecule / structure_id", "Region development", "可检查页面/crop 回归"],
        ["wild_strict_v3", "301", "paper_group + molecule", "Locked final", "冻结后一次性运行，不回流"],
        ["scaffold_novel", "134", "Bemis-Murcko scaffold", "Locked 子集", "诊断训练未见骨架泛化"],
        ["wild_symbolic_v3", "460", "symbolic label", "独立文字轨道", "不做 canonicalization，不并主分数"],
        ["private_photo_v3", "0", "structure_id", "计划 locked", "当前不得声称已完成自采"],
    ], widths=[4.2, 1.3, 4, 3.3, 4.1])
    add_heading(doc, "5.1 Development 配比的讲究", level=2)
    add_paragraph(doc, "Core 预设 DECIMER/UOB/USPTO/real-world 配额为 150/200/200/217，严格 QC 后为 150/193/196/214，共 753；region 预设 EDU-CHEMC/UOB/USPTO/real-world 为 153/200/200/217，严格后为 151/193/196/214，共 754。UOB/USPTO 各约 200 是可复现 printed 锚点，不代表部署自然频率；DECIMER/EDU 约 150 让手绘/教育困难样本有统计可见度；real-world 清洗后全部保留，避免稀缺场景被下采样抹掉。")
    add_heading(doc, "5.2 论文级 Locked 构建", level=2)
    add_paragraph(doc, "MolRecBench 原始 5,008 条中，3,508 条为 symbolic/非法 canonical 标签，72 条与 legacy development 分子重叠；剩余 strict pool 为 1,428 条、来自 519 篇论文。构建器按 paper_group 留出 62 篇完整论文，每篇最多 5 图，形成 301 张图、301 个唯一 canonical 分子。留出论文中的另外 308 张继续 held out，不回流训练；跨论文但 canonical 与 locked 重复的 19 条从训练侧删除。")
    add_callout(doc, "防泄漏核心", "同一论文的字体、版式、裁图风格高度相关；按图片随机切分会把这些视觉指纹泄漏到训练。按 paper_group 整篇留出，比行级随机切分更接近真正的论文外推。", color=RED)
    add_heading(doc, "5.3 人工审核与锁定", level=2)
    add_paragraph(doc, "项目 owner 已声明 wild strict、symbolic 和两个 legacy development 冻结 labels 完成离线人工审核，且审核后没有剔除或修改，因此现有指标无需重算。声明由四个 labels 的 SHA256 绑定；文件变化时声明自动失效。该证据不是独立双人盲审，也不虚构审核姓名、签名、分歧数量或逐样本决定。")

    add_heading(doc, "6. 训练与消融设计", level=1, new_page=True)
    add_paragraph(doc, "所有主因子对照使用同一 V2-1 基座、相同 LoRA 结构、相同步数和 evaluation 面板。2×2 因子为 strict wild 训练数据是否加入与离线退化是否加入，每个条件运行两个 seed；另做 warm-start 和增强剂量诊断。核心统计量为两 seed 宏平均 exact、seed 范围和最低 valid rate。")
    add_table(doc, ["条件", "Wild", "Aug", "两 seed mean macro exact", "解释"], [
        ["00 control", "Off", "Off", "34.11%", "本轮探索均值最高，选择为 final 数据"],
        ["10 wild-only", "On", "Off", "33.28%", "单独加入未稳定提升"],
        ["01 aug-only", "Off", "On", "33.44%", "单独加入未稳定提升"],
        ["11 wild + aug", "On", "On", "33.91%", "存在正交互迹象，但不足以胜过 control"],
    ], widths=[3, 2, 2, 4.2, 6.1])
    add_paragraph(doc, "两个 seed 仍不足以估计训练随机性的尾部，运行顺序也没有完全随机化或位置平衡。因此 V3 的正确结论是“在本轮预算和 development 面板下选择 control”，而不是“真实数据或增强无效”。paired bootstrap 用于同一 development 样本上的模型差异，不可替代独立训练 seed。")
    add_heading(doc, "6.1 Final continuation 参数", level=2)
    add_table(doc, ["项目", "设置", "理由"], [
        ["基座", "V2-1 export", "保留已学会的 OCSR 输出格式"],
        ["训练数据", "A control 22,762", "2×2 探索均值最高且 validity 过门"],
        ["训练长度", "1,400 steps", "development checkpoint 选择"],
        ["参数高效微调", "LoRA continuation", "降低显存和全量权重漂移"],
        ["选模面板", "753 + 754 development", "同时守住 core 和 region"],
        ["主指标", "macro canonical exact", "避免一个面板支配决策"],
        ["安全门", "validity floor + 0.5pp gate", "复杂策略必须有足够收益且不回归"],
    ], widths=[4.2, 5.1, 7.2])

    add_heading(doc, "7. 后训练与解码选择", level=1, new_page=True)
    add_paragraph(doc, "V3 把训练后决策分成三类：权重继续训练、生成搜索和候选重排。所有策略先在 development 上按预设闸门比较，只有模型、prompt、checkpoint、decoder 和文件 hash 全部冻结后，才运行 locked test。")
    add_table(doc, ["候选策略", "Development macro exact", "相对基准", "决策", "为什么"], [
        ["ckpt-1400 greedy", "35.97%", "基准", "保留", "最终 checkpoint"],
        ["300-step hard replay", "35.24%", "-0.73pp", "拒绝", "未达 +0.5pp 门且回归"],
        ["beam4 / return4", "42.07%", "+6.10pp", "采用", "两面板方向一致，收益明显"],
        ["chem-light rerank", "39.55%", "-2.52pp", "拒绝", "同候选池上劣于原始 beam 排序"],
    ], widths=[4.2, 3.5, 2.5, 2, 5.2])
    add_paragraph(doc, "Beam 搜索有效的原因是 OCSR 输出中局部字符歧义会产生多个接近候选；单次 greedy 过早提交，而 beam4/return4 保留更多全局序列。RDKit 有效性、canonical 投票和生成分数用于候选证据解释，但额外 chem-light 排序在当前候选池上降低 exact，因此没有以“化学规则更复杂”为理由强行采用。")
    add_callout(doc, "后训练为什么这么搞", "复杂方法只有在同候选池、同面板、预设收益门下胜出才进入 final。Hard replay 和 rerank 被拒绝本身是重要证据：项目没有只展示成功实验，也没有把复杂度当创新。", color=BLUE)

    add_heading(doc, "8. 最终结果与误差边界", level=1, new_page=True)
    add_table(doc, ["面板", "N", "Exact", "Valid/Nonempty", "解释"], [
        ["legacy core development", "753", "41.70%", "—", "beam4/return4 选模面板之一"],
        ["legacy region development", "754", "42.44%", "—", "beam4/return4 选模面板之一"],
        ["wild strict locked", "301", "22.92%", "84.72%", "62 篇完整论文外推"],
        ["scaffold novel locked", "134", "13.43%", "75.37%", "训练未见骨架诊断"],
        ["symbolic independent", "460", "0.00%", "100.00%", "不同任务定义，不并主分数"],
    ], widths=[4.4, 1.5, 2.3, 3.3, 5.8])
    add_paragraph(doc, "Wild exact 明显低于 development，而 valid 仍达到 84.72%，说明主要错误不是完全无法生成 SMILES，而是键级、立体化学、取代位置或复杂拓扑的细微错误。Scaffold novel 进一步下降到 13.43%，支持“未见骨架 + 论文域视觉变化”是当前最难组合的判断。")
    add_paragraph(doc, "评测主指标为 RDKit canonical exact。Valid SMILES、token F1、normalized edit similarity 和 fingerprint Tanimoto 用于解释错误，不替代 exact；否则一个化学上有效但结构错误的候选会被误算为成功。Symbolic 的 0% 只说明当前 canonical decoder 不适合 R-group/缩写转写，不能反推 canonical OCSR 无效。")
    add_heading(doc, "8.1 最可能的错误来源", level=2)
    add_table(doc, ["错误层", "典型现象", "诊断指标", "下一步"], [
        ["视觉域移", "论文截图、低分辨率、文本邻近", "wild vs development gap", "真实域训练与版式增强"],
        ["键级/环闭合", "SMILES 有效但 exact 错", "valid 高、exact 低", "候选级图结构一致性模型"],
        ["立体化学", "@/@@、楔线方向错误", "stereo-stripped exact", "立体专门数据与多尺度 crop"],
        ["长分子", "token 截断或局部错位", "长度分层、edit similarity", "长度感知 curriculum/解码"],
        ["symbolic", "R-group、缩写无法 canonicalize", "独立转写 exact", "专用 prompt/decoder/标签协议"],
    ], widths=[3, 5.2, 3.8, 5.2])

    add_heading(doc, "9. MolTrace Agent 前后端", level=1, new_page=True)
    add_picture(doc, figures["agent"], "图 4  Agent 将黑盒生成拆成六步可审计决策")
    add_picture(doc, ASSET_DIR / "agent-workbench-desktop.png", "图 5  真实浏览器验收截图：桌面端完整 OCSR 工作台", width=6.9)
    add_paragraph(doc, "前端提供拖放/相机来源图片、预览、亮度/对比度/尺寸质量门、beam/return/max-token/TTA 控制、六步 trace、canonical 结果、候选表、运行历史和 JSON 导出。界面采用克制的工作台布局，不把识别任务包装成无关的聊天窗口。")
    add_paragraph(doc, "后端采用 Node.js 18+ 标准库，无 npm 运行依赖。设置 V3_MODEL_DIR 与 PYTHON_BIN 后，固定调用现有 infer_ocsr_transformers.py，结果继续经过 RDKit canonicalization 和候选排序；没有 GPU 模型时，只允许对内置咖啡因已知标签展示引导流程，任意用户上传图会返回 needs_model，绝不伪造识别结果。")
    add_table(doc, ["API", "作用", "关键边界"], [
        ["GET /api/health", "服务、模型和历史状态", "用于部署探针，不证明模型精度"],
        ["GET /api/model", "冻结 checkpoint/decoder/指标/revision", "数字来自 evidence，不在线重算"],
        ["POST /api/agent/run", "六步识别、候选和 trace", "模型不可用时拒绝任意图伪预测"],
        ["POST /api/validate", "轻量 SMILES 词法预检", "明确不替代 RDKit 化学解析"],
        ["GET/DELETE /api/history", "最近 20 条摘要", "不持久化原图，仅指纹和结果摘要"],
    ], widths=[4.6, 5.3, 7.2])
    add_callout(doc, "应用验收", "Node 单元测试 5/5 通过；真实浏览器在 1440×1050 和 390×844 视口运行完整示例，四个候选和六步 trace 正常渲染，横向溢出为 0。截图与 browser-audit.json 随应用保存。")

    add_heading(doc, "10. 许可、复现与交付", level=1, new_page=True)
    add_paragraph(doc, "项目代码、文档和可发布派生权重按 Apache-2.0/NOTICE 发布。数据集名称不是再分发许可；训练 manifest 的部分历史记录只有集合级 source，样本级 license、source URL 和 structure_id 尚未全覆盖，因此公共 GitHub 不重新分发训练原图或完整私有 JSONL。评测资产按各上游许可和来源文档管理，公开仓保留构建方法、QC、统计和 hash 证据。")
    add_table(doc, ["对象", "公开策略", "证据/入口"], [
        ["源代码与脚本", "GitHub 公开", "LICENSE、NOTICE、commit 固定"],
        ["V3 模型", "Hugging Face + 本地 ZIP", "revision 与 SHA256 manifest"],
        ["训练原图/受限数据", "不在 GitHub 重分发", "DATA_LICENSES_AND_ATTRIBUTION_zh.md"],
        ["评测构建方法", "报告、README、QC、builder 公开", "EVAL_DATASET_CONSTRUCTION_REPORT_zh.md"],
        ["人工审核", "公开 owner attestation + labels hash", "不虚构审阅者身份和逐样本记录"],
        ["Agent 应用", "前后端源码、API 与测试公开", "V3/agent_demo/"],
    ], widths=[4.2, 5.3, 7.6])
    add_heading(doc, "10.1 决赛交付清单", level=2)
    add_table(doc, ["材料", "最终文件", "验收重点"], [
        ["总报告", "数据构建报告_V3_final.docx", "训练、评测、方法、Agent、许可齐全"],
        ["评测报告", "评测集数据构建报告_V3_final.docx", "目录、数量、配比、防泄漏、QC、边界"],
        ["答辩稿", "答辩PPT_V3_final.pptx + HTML", "18 页、可编辑文字、真实界面、逐页审计"],
        ["评测包", "V3_evaluation_dataset_complete_20260719.zip", "独立 README、manifest 与警告边界"],
        ["模型包", "paddleocr_weights_V3_final.zip", "loadable 模型目录 + README + SHA256SUMS"],
        ["开源项目", "GitHub + Hugging Face", "commit/revision 固定、链接可访问"],
    ], widths=[3.1, 7.3, 6.7])

    add_heading(doc, "11. 限制与后续计划", level=1, new_page=True)
    add_paragraph(doc, "当前 private-photo locked test 为 0，不能把算法退化图或公开论文图写成自采；两 seed 2×2 是探索性实验，不足以支持统计显著最优；owner attestation 不是独立双人盲审；第二台机器 clean-download 和容器 digest 复现尚未完成；样本级 license/source URL/structure_id 覆盖不足，限制了训练数据公开再分发。")
    add_table(doc, ["风险", "当前控制", "下一步可证伪动作"], [
        ["训练 seed 偶然性", "两个 seed + seed range", "至少四 seed confirmatory，平衡运行顺序"],
        ["部署域偏移", "论文级 locked wild", "新增授权真实拍照，多设备/光照分组"],
        ["审核独立性", "owner attestation + hash", "双人盲审、一致性统计、仲裁记录"],
        ["数据许可", "不重新分发未知来源原图", "样本级 license/source URL 补全与隔离"],
        ["应用性能", "CLI 子进程失败隔离", "模型常驻服务、任务队列、并发与显存监控"],
        ["Symbolic 失败", "独立轨道诚实报告", "专用标签协议、prompt 和 decoder development"],
    ], widths=[3.8, 6.2, 7.1])
    add_callout(doc, "最终表述", "本项目证明了从 V2-1 task adaptation 到 V3 数据治理、受控 continuation、decoder 选择、一次性 locked test 和可审计 Agent 的完整流程。它没有证明 OCSR 问题已经解决；反而用 22.92% wild exact 与 13.43% scaffold exact 精确界定了下一阶段需要解决的真实域和骨架泛化问题。", color=TEAL)

    add_heading(doc, "附录 A. 证据索引", level=1, new_page=True)
    add_table(doc, ["证据", "路径/版本"], [
        ["最终结果", "V3/evidence/FINAL_RESULTS.json 与 FINAL_RESULTS_zh.md"],
        ["数据构建", "V3/evidence/dataset_build_report.json"],
        ["混合统计", "V3/evidence/mixture_counts.csv"],
        ["选 checkpoint", "V3/evidence/final_checkpoint_selection.json"],
        ["Hard replay", "V3/evidence/final_vs_hard_replay.json"],
        ["Decoder", "V3/evidence/generation_policy_beam_selection.json"],
        ["人工审核", "V3/qc/MANUAL_REVIEW_ATTESTATION_zh.md + JSON"],
        ["评测方法", "V3/runbooks/EVALUATION_PROTOCOL_zh.md"],
        ["Agent", "V3/agent_demo/README.md + output/playwright/browser-audit.json"],
        ["GitHub", "https://github.com/2658183739/-PaddleOCR-VL-1.5-OCSR"],
        ["Hugging Face", "https://huggingface.co/L2658183739/PaddleOCR-VL-1.5-OCSR"],
    ], widths=[4.1, 13])
    add_heading(doc, "附录 B. 方法参考", level=1)
    references = [
        "Hu, E. J. et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.",
        "Bemis, G. W. & Murcko, M. A. The Properties of Known Drugs. 1. Molecular Frameworks. J. Med. Chem., 1996.",
        "Efron, B. & Tibshirani, R. J. An Introduction to the Bootstrap. Chapman & Hall, 1993.",
        "RDKit: Open-source cheminformatics software. Canonical SMILES and molecular validation documentation.",
        "DECIMER / MolScribe / MolGrapher / MolRecBench publications and their associated dataset documentation, used according to source-specific attribution and redistribution terms.",
    ]
    for reference in references:
        add_paragraph(doc, reference, size=8.5, color=MUTED, space_after=4)

    output = FINAL_DIR / "数据构建报告_V3_final.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def build_eval_report(figures):
    doc = Document()
    configure_document(doc, "评测集数据构建报告")
    add_cover(doc, "PaddleOCR-VL OCSR V3\n评测集数据构建报告", "从 V2-1 历史面板到论文级锁定终测的角色、配比、质量控制与证据边界", "FINAL EVALUATION DATASET REPORT")
    sections = [
        ("V2-1 与 V3 的角色迁移", "为什么历史面板必须降级为 development"),
        ("评测目录与数量", "每个目录装了什么"),
        ("Development 配比", "锚点与稀缺场景如何平衡"),
        ("Locked wild 构建", "如何实现论文外推"),
        ("Scaffold 与 symbolic", "主任务之外如何独立诊断"),
        ("质量控制与人工审核", "自动检查和人工声明各自证明什么"),
        ("指标、结果与锁定协议", "哪些分数可比较、哪些不能"),
        ("许可、打包与复核", "评测包如何交付"),
    ]
    add_toc(doc, sections)
    add_heading(doc, "摘要", level=1)
    add_paragraph(doc, "V3 将曾被反复用于调参的 V2-1 历史主面板明确降级为 legacy development，并新增按论文整组留出的 wild strict final test。最终 development 为 753 + 754 条；locked wild 为 301 张图、301 个唯一 canonical 分子、来自 62 篇完整留出论文；其中 134 条为训练未见 Bemis-Murcko scaffold；460 条 R-group/缩写记录进入独立 symbolic 轨道。")
    add_callout(doc, "评测原则", "先冻结数据角色，再训练和选模；development 可以比较策略，locked test 只在全部策略冻结后运行一次；symbolic 不属于单分子 canonical SMILES，必须单独报告。")
    add_picture(doc, figures["eval"], "图 1  V3 评测集按角色拆分")

    add_heading(doc, "1. V2-1 与 V3 的角色迁移", level=1, new_page=True)
    add_paragraph(doc, "V2-1 的 canonical_smiles_main_v1（767 条）和 region_panel_770（770 条）已经被用于历史训练诊断与后处理选择，因此不再具备“未触碰测试集”的含义。V3 对它们做严格 canonical 清洗，形成 753 条 core development 和 754 条 region development。V2-1 core 37.03% → 45.89%、region 38.44% → 43.77% 证明历史改进，但不能与 V3 locked wild 22.92% 混算提升率。")
    add_table(doc, ["版本", "面板", "N", "角色", "可做的决策"], [
        ["V2-1", "canonical_smiles_main_v1", "767", "历史诊断", "回溯 task adaptation 和后处理"],
        ["V2-1", "region_panel_770", "770", "历史诊断", "回溯页面/区域方案"],
        ["V3", "dev_legacy_core_strict", "753", "Development", "数据因素、checkpoint、decoder"],
        ["V3", "dev_legacy_region_strict", "754", "Development", "页面/crop 回归"],
        ["V3", "wild_strict_v3", "301", "Locked final", "冻结后一次性外推报告"],
    ], widths=[2.2, 5.1, 1.4, 3.2, 5.2])

    add_heading(doc, "2. 评测目录与数量", level=1, new_page=True)
    add_table(doc, ["目录", "N", "主要内容", "V3 角色"], [
        ["canonical_smiles_main_v1/", "767", "labels JSONL/CSV + 四来源图像", "V2-1 历史；派生 core dev"],
        ["ocsr_realworld_mixed_eval_v1p1/", "770", "labels、images、来源/QC/技术报告", "V2-1 历史；派生 region dev"],
        ["dev_legacy_core_strict/", "753", "冻结 labels；相对路径引用图像", "Core development"],
        ["dev_legacy_region_strict/", "754", "教育/printed/real-world 严格 labels", "Region development"],
        ["molrecbench_wild_300/", "300", "历史辅助 wild 子集", "不是 V3 locked final"],
        ["wild_strict_v3/", "301", "论文级留出 labels", "Locked canonical final"],
        ["wild_strict_scaffold_novel_v3/", "134", "wild 的未见 scaffold 子集", "Locked 泛化诊断"],
        ["wild_symbolic_v3/", "460", "R-group/缩写原始转写", "独立 symbolic track"],
    ], widths=[5.3, 1.3, 6.1, 4.4])
    add_paragraph(doc, "完整评测包保留 README、目录清单、labels、资产引用、QC、来源与哈希说明；公开 GitHub 只放可公开方法和证据，不无条件重新分发受限图片。最终压缩包为 V3_evaluation_dataset_complete_20260719.zip。")

    add_heading(doc, "3. Development 配比", level=1, new_page=True)
    add_table(doc, ["面板", "目标配额", "Strict 通过", "科学目的"], [
        ["Core", "DECIMER 150 / UOB 200 / USPTO 200 / real-world 217", "150 / 193 / 196 / 214 = 753", "printed 锚点 + 手绘 + 稀缺真实场景"],
        ["Region", "EDU 153 / UOB 200 / USPTO 200 / real-world 217", "151 / 193 / 196 / 214 = 754", "教育图/页面区域 + printed 锚点 + 真实退化"],
    ], widths=[2.6, 6.7, 4.6, 3.9])
    add_paragraph(doc, "评测集不是把所有上游数据按自然数量混合。UOB/USPTO 各约 200 提供跨版本可比锚点；DECIMER/EDU 约 150 保证高风险视觉类型不会被总体平均掩盖；real-world 清洗后全部保留，避免稀缺场景因随机下采样消失。严格通过数少于目标配额时不回填，差异必须由非法标签、坏图、任务边界或泄漏记录解释。")
    add_callout(doc, "配比不是部署先验", "这些比例服务于受控比较和短板发现，不代表真实部署流量。若部署到专利扫描、手写实验记录或论文截图，应另建按目标场景分层的外部测试。", color=BLUE)

    add_heading(doc, "4. Locked wild 构建", level=1, new_page=True)
    add_paragraph(doc, "MolRecBench 原始 5,008 条经过任务与标签筛选：3,508 条 symbolic/非法 canonical 记录不进入主任务，72 条与 legacy development 分子重叠；剩余 strict pool 为 1,428 条、来自 519 篇论文。构建器按 paper_group 留出 62 篇完整论文，每篇最多 5 图，最终得到 301 图、301 个唯一 canonical 分子。")
    add_table(doc, ["步骤", "输入", "输出/拒绝", "目的"], [
        ["任务筛选", "5,008", "3,508 symbolic/非法分流", "固定 canonical 主任务"],
        ["Dev 去重", "strict 候选", "72 molecule overlap 排除", "避免历史面板泄漏"],
        ["Strict pool", "剩余 canonical", "1,428 / 519 papers", "论文级抽样母体"],
        ["整篇留出", "519 papers", "62 evaluation papers", "阻断论文版式/字体泄漏"],
        ["每篇上限", "留出论文", "最多 5 图/论文", "避免大论文主导总分"],
        ["最终锁定", "候选图", "301 图/301 分子", "一次性 final test"],
        ["训练反向删除", "跨论文记录", "19 locked molecule overlap", "保证 molecule 零重叠"],
    ], widths=[3.1, 3.3, 5.5, 5])
    add_paragraph(doc, "同一论文中的字体、图像压缩、标注风格和裁图方式相似，行级随机切分会把这些视觉指纹同时放入训练和测试。paper_group 整篇留出直接回答“遇到新论文时是否泛化”，是比简单随机抽样更保守也更符合部署的问题。")

    add_heading(doc, "5. Scaffold 与 symbolic", level=1, new_page=True)
    add_paragraph(doc, "在 301 条 wild strict 中，根据 Bemis-Murcko scaffold 与训练覆盖比较，得到 134 条训练未见 scaffold 子集。它共享同一次 locked 评测，不能单独调参；exact 13.43%、valid 75.37% 表明新骨架和论文域移叠加后仍明显困难。")
    add_paragraph(doc, "MolRecBench 中的 R-group、缩写和 symbolic 标签不属于单分子 canonical SMILES。460 条记录保留原始转写口径并使用 whitespace-normalized exact 与 nonempty rate 独立报告。当前 exact 0%、nonempty 100%，说明模型会输出内容但输出协议不匹配；下一步应使用专用标签规则、prompt 和 decoder development，而不是把它们混进主任务稀释结论。")
    add_table(doc, ["轨道", "标签空间", "主指标", "是否参与 canonical 选模"], [
        ["Wild strict", "单分子 canonical SMILES", "RDKit canonical exact", "否；只在冻结后报告"],
        ["Scaffold novel", "同 wild strict", "canonical exact + valid", "否；派生诊断"],
        ["Symbolic", "R-group/缩写文字", "whitespace-normalized exact", "否；独立 decoder 才可选模"],
    ], widths=[3.3, 5.5, 4.4, 4])

    add_heading(doc, "6. 质量控制与人工审核", level=1, new_page=True)
    add_table(doc, ["QC 层", "检查项", "证据", "局限"], [
        ["结构自动 QC", "字段、路径、图片、标签类型、重复", "dataset_build_report / QC report", "不能判断每个键级标签是否正确"],
        ["化学自动 QC", "RDKit parse、canonical、片段、dummy", "builder 与 evaluation scripts", "symbolic 不适用"],
        ["泄漏 QC", "molecule、structure_id、paper_group", "split manifest 与 overlap counts", "依赖标识字段质量"],
        ["人工复核", "图像-标签一致性、任务边界", "owner attestation + 4 labels hash", "不是独立双人盲审"],
        ["锁定 QC", "labels hash、策略 hash、禁止回流", "runbook + public release evidence", "第二台机器复现仍待补"],
    ], widths=[3.3, 5.1, 5.2, 3.9])
    add_paragraph(doc, "项目 owner 已确认 wild strict、symbolic、core dev 和 region dev 的冻结 labels 完成离线审核，审核后没有删除或修改，因此结果不需要重算。公开 JSON 将声明绑定到四个文件 SHA256，任何文件变化都会使声明失效。为保持证据诚实，公开材料不虚构 reviewer 姓名、签名、双盲分歧数或逐样本审阅记录。")
    add_callout(doc, "自动 QC 不等于人工 QC", "前者证明格式、路径、化学语法和泄漏规则；后者才检查图像与结构标签是否对应。当前证据是 owner-attested complete，不应写成独立第三方审计。", color=AMBER)

    add_heading(doc, "7. 指标、结果与锁定协议", level=1, new_page=True)
    add_picture(doc, figures["results"], "图 2  历史、development 与 locked 的结果边界")
    add_table(doc, ["指标", "回答的问题", "主/辅", "常见误用"], [
        ["Canonical exact", "结构是否完全一致", "主指标", "把不同面板混算提升率"],
        ["Valid SMILES", "字符串能否表示合法分子", "辅助", "把合法但错误结构算成功"],
        ["Token F1 / edit", "错误是局部还是整体", "辅助", "替代结构 exact"],
        ["Fingerprint Tanimoto", "拓扑相似程度", "辅助", "高相似即宣称识别正确"],
        ["Scaffold exact", "未见骨架泛化", "诊断", "用子集反向调参"],
        ["Symbolic exact", "原始文字协议是否匹配", "独立主指标", "混入 canonical 分母"],
    ], widths=[3.7, 5.4, 2.3, 5.1])
    add_paragraph(doc, "Development 只选择 checkpoint-1400、拒绝 hard replay、采用 beam4/return4 并拒绝 chem-light rerank。锁定策略后运行 wild strict：exact 22.92%、valid 84.72%；scaffold novel exact 13.43%、valid 75.37%。这些结果不回流训练、prompt、beam 或候选排序。")

    add_heading(doc, "8. 许可、打包与复核", level=1, new_page=True)
    add_paragraph(doc, "评测清单保留 source、上游 ID 和来源文档。项目代码与可发布派生权重采用 Apache-2.0/NOTICE，但上游数据仍按各自许可；公共仓不无条件重新分发受限图片和逐样本预测。评测压缩包用于比赛本地复核，公开分发前仍应逐项检查许可与隐私边界。")
    add_table(doc, ["复核对象", "文件/目录", "验收动作"], [
        ["目录说明", "V3/data/eval/README.md", "核对历史/开发/locked/symbolic 标记"],
        ["构建统计", "V3/evidence/dataset_build_report.json", "核对输入、过滤、重叠和最终计数"],
        ["冻结结果", "V3/evidence/FINAL_RESULTS.json", "核对模型/decoder/locked 指标"],
        ["人工声明", "V3/qc/manual_review_attestation.json", "核对四个 labels SHA256"],
        ["评测协议", "V3/runbooks/EVALUATION_PROTOCOL_zh.md", "核对禁止回流和统计单位"],
        ["提交包", "V3_evaluation_dataset_complete_20260719.zip", "解压、README、manifest 与 hash"],
    ], widths=[4.1, 7.1, 5.8])
    add_callout(doc, "提交声明", "当前没有 private-photo locked test；owner attestation 不是独立复现；任何 locked test 后的继续调参都只能标记 exploratory。报告据此不夸大评测集覆盖和外推结论。")

    output = FINAL_DIR / "评测集数据构建报告_V3_final.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def structural_audit(paths):
    audit = {}
    for path in paths:
        doc = Document(path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        audit[path.name] = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "sections": len(doc.sections),
            "headings": headings,
            "bytes": path.stat().st_size,
        }
    audit_path = FINAL_DIR / "submission_documents_audit.json"
    audit_path.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    return audit_path


def main():
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    figures = build_figures()
    main_report = build_main_report(figures)
    eval_report = build_eval_report(figures)
    audit = structural_audit([main_report, eval_report])
    print(main_report)
    print(eval_report)
    print(audit)


if __name__ == "__main__":
    main()
